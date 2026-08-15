"""문서 파싱 + 청킹 모듈 (비ML 추출).

 형식별 섹션 파서 → _pack_sections(묶기) → SentenceSplitter(크기 cap) → ChunkData.
 워커(rag/documents.py)가 이 결과를 받아 임베딩 + DB 적재.

 **형식 분기는 chunk_file 한 곳뿐이다** — 두 곳에 있던 게 CLI로 xlsx를 넣으면 ZIP
 바이너리가 색인되던 버그의 원인이었다(#42).

 ================================================================
   형식별 파싱 — heading_path는 xlsx 포함 네 형식 모두 채운다 (실측 100%, 2026-08-15)
 ================================================================
 - PDF : pdfplumber로 줄별 글자 크기를 보고 헤딩 판정. page(첫 본문 줄 쪽)도 보존 — 유일.
 - DOCX: python-docx로 body를 문서 순서대로 읽어 Heading 스타일 기준 섹션 분할.
 - MD  : '#' 헤딩 정규식 + 코드 펜스 가드.
 - TXT : 평문(인코딩 감지 utf-8→cp949). 헤딩 구조가 없어 heading_path는 빈다.
 - XLSX: rag/xlsx_chunking (openpyxl). 시트=1청크라 이 파이프라인 밖 — 분할하지 않고
         150행 초과면 업로드를 거절한다.

 공통: SentenceSplitter(chunk_size=512, chunk_overlap=50)로 크기 cap.
   - 큰 섹션이 비대해져 검색 정확도 떨어지는 것 방지 + LLM 컨텍스트 부담 완화
   - overlap 50으로 경계 의미 단절 완화

 결과 ChunkData: text / heading_path / page / chunk_index / meta.
 (LlamaIndex 의존성은 이 모듈 안에 격리 — 호출부는 ChunkData 도메인 타입만 본다.)
 """

import re
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from llama_index.core import Document as LiDocument
from llama_index.core.node_parser import MarkdownNodeParser, SentenceSplitter

_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$')
_DOCX_HEADING_RE = re.compile(r'^Heading (\d+)$')     # python-docx는 빌트인 헤딩을 영문명으로 준다
# 섹션 묶음 예산(문자). 섹션 하나하나는 보통 100~150자라, 섹션=청크로 두면 top_k=5가
# 실어 나르는 컨텍스트가 옛 방식의 1/3로 줄어든다(실측 2211자→707자).
# 주의: 이 값이 상한으로 작동하는 일은 드물다 — _pack_sections의 경계 규칙(같은 대분류
# 아래 형제끼리만)이 먼저 걸려 실측 청크 중앙값은 200자대다. 실제 토큰 상한은 SentenceSplitter.
_PACK_CHARS = 500

@dataclass
class ChunkData:
    """청킹 단계 중간 결과 (임베딩 / DB insert 전)."""
    text: str                   # 청크 본문
    heading_path: list[str]     # 예: ["3. 배송지연 보상", "3.2 지급 기준"]
    page: int | None            # PDF는 페이지 번호, DOCX는 None
    chunk_index: int            # 문서 내 순서 (0부터)
    meta: dict | None = None    # F1a: 청크 metadata (xlsx는 {"is_table": True, "sheet": 시트명})


@dataclass
class _Section:
    """형식별 파서 → _pack_sections 사이의 내부 계약. 이 모듈 밖으로 나가지 않는다.

    파서 넷(_pdf/_docx/_md/_txt_sections)이 전부 이 형태를 내놓기 때문에 그 뒤 단계
    (묶기 → 크기 분할 → ChunkData)를 형식과 무관하게 한 번만 쓴다.
    """
    heading_path: list[str]
    page: int | None            # PDF만 채운다 — 나머지는 파일에 페이지 개념이 없다
    body: str

def _read_text(file_path) -> str:
    """텍스트 파일을 인코딩 감지해 읽는다 (P2 CP949).
    utf-8 우선 → 실패 시 cp949(국내 txt 빈발) 폴백 → 그래도 실패면 replace 최후.
    (기존 utf-8+errors='replace'는 CP949를 조용히 �로 깨뜨렸음.)
    """
    raw = Path(file_path).read_bytes()
    for enc in ('utf-8', 'cp949'):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode('utf-8', errors='replace')


def _extract_heading_path(metadata: dict, text: str) -> list[str]:
    """조상 헤딩(metadata.header_path) + 자기 헤딩(text 첫 줄)을 합쳐 반환.

    MarkdownNodeParser는 'header_path'에 조상 경로만 '/' 구분자로 저장하고,
    자기 자신의 헤딩은 chunk text 첫 줄에 그대로 둠.

    예: metadata={'header_path': '/3. 표준 상담 스크립트/'}, text='### 3.1 ...\\n...'
        -> ['3. 표준 상담 스크립트', '3.1 ...']
    """
    path_str = metadata.get('header_path', '/')
    ancestors = [seg for seg in path_str.split('/') if seg]

    # 본문 첫 줄에서 현재 헤딩 추출 (있으면)
    first_line = text.lstrip().split('\n', 1)[0]
    m = _HEADING_RE.match(first_line)
    if m:
        ancestors.append(m.group(2).strip())

    return ancestors

def _pdf_pages(file_path: str | Path) -> list[tuple[int, str]]:
    """PDF를 페이지별 텍스트로 (page 번호 보존 — 인용용). 텍스트 레이어 없으면 빈 문자열.
    비ML 추출(pdfplumber) — 표는 행/열이 공백 구분 텍스트로 나오며 구조 서식은 없다.
    """
    with pdfplumber.open(str(file_path)) as pdf:
        return [(i, page.extract_text() or '') for i, page in enumerate(pdf.pages, start=1)]


def _pdf_lines(file_path: str | Path) -> list[tuple[int, str, float]]:
    """PDF를 (page, 줄 텍스트, 최대 글자크기) 목록으로. 헤딩 판정에 크기가 필요해서.

    extract_text()는 크기 정보를 버리므로 extract_text_lines()로 줄별 char를 본다.
    """
    out: list[tuple[int, str, float]] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            for line in page.extract_text_lines():
                text = (line.get('text') or '').strip()
                if not text:
                    continue
                sizes = [c['size'] for c in line.get('chars', []) if c.get('size')]
                out.append((page_no, text, round(max(sizes), 1) if sizes else 0.0))
    return out


def _pdf_sections(file_path: str | Path) -> list[_Section]:
    """PDF를 (heading_path, 시작 page, 본문) 섹션 목록으로 — 글자 크기로 헤딩을 판정.

    DOCX와 달리 PDF엔 구조 태그가 없어 조판에 기대는 휴리스틱이다:
      본문 크기 = 가장 많이 쓰인 글자 크기, 그보다 **큰** 줄을 헤딩으로 본다.
      크기 종류를 큰 순으로 정렬해 1,2,3… 레벨로 매긴다.
    크기 차이가 없는 문서(전부 같은 크기, 스캔본 등)는 헤딩 0개 → heading_path=[]로
    폴백해 이 변경 전과 동일하게 동작한다 (조판이 달라도 손해는 없게).
    """
    lines = _pdf_lines(file_path)
    if not lines:
        return []

    # 본문 크기 = 글자 수 기준 최빈 크기 (줄 수가 아니라 분량 기준이라 제목에 안 휘둘린다)
    weight: dict[float, int] = {}
    for _, text, size in lines:
        weight[size] = weight.get(size, 0) + len(text)
    body_size = max(weight, key=lambda s: weight[s])

    # 본문보다 큰 크기들 → 큰 순으로 레벨 부여
    level_of = {size: lv for lv, size in
                enumerate(sorted((s for s in weight if s > body_size), reverse=True), start=1)}

    sections: list[_Section] = []
    stack: list[str] = []
    buf: list[str] = []
    start_page = lines[0][0]

    def flush() -> None:
        if buf:              # 본문 없는 헤딩은 섹션이 되지 않는다 (docx·md도 같은 가드)
            sections.append(_Section(list(stack), start_page, '\n'.join(buf)))
            buf.clear()

    for page_no, text, size in lines:
        level = level_of.get(size)
        if level:
            flush()
            del stack[level - 1:]
            stack.append(text)
        else:
            # 섹션의 page는 '첫 본문 줄'의 페이지다 — 헤딩이 페이지 끝에 걸리면
            # 본문은 다음 쪽에서 시작하고, 인용은 본문이 있는 쪽을 가리켜야 맞다.
            if not buf:
                start_page = page_no
            buf.append(text)
    flush()
    return sections


def _docx_items(file_path: str | Path):
    """DOCX body를 **문서 원래 순서대로** (is_heading, level, text)로 흘린다.

    doc.paragraphs와 doc.tables를 따로 훑으면 표가 전부 문서 끝으로 밀린다
    (예: '2.1 기준표' 밑의 표가 자기 섹션에서 떨어져 나와 마지막 청크로 감).
    body 자식을 직접 순회해 문단·표의 상대 순서를 보존한다.
    """
    doc = DocxDocument(str(file_path))
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            para = DocxParagraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            # style이 None인 문단이 실제로 있다 (스타일 정의가 빠진 docx) — getattr로 방어
            m = _DOCX_HEADING_RE.match(getattr(para.style, 'name', None) or '')
            yield (bool(m), int(m.group(1)) if m else 0, text)
        elif tag == 'tbl':                         # docx 표는 구조화 포맷 → 셀 값 안정 추출
            for row in DocxTable(child, doc).rows:
                # 빈 셀도 자리를 지킨다 — 걸러내면 뒷 컬럼이 앞으로 당겨져 헤더와 어긋난다.
                # ['제주','','5000원','도서산간 별도'] → '제주 | 5000원 | 도서산간 별도'가 되어
                # LLM이 '조건=5000원'으로 읽는다. 완전 빈 행만 스킵 (xlsx 파서와 같은 원리).
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    yield (False, 0, ' | '.join(cells))


def _docx_text(file_path: str | Path) -> str:
    """DOCX 전체를 평문 하나로 (채팅 첨부 — 컨텍스트 직접 주입용). 헤딩 포함, 문서 순서 보존."""
    return '\n'.join(text for _, _, text in _docx_items(file_path))


def _docx_sections(file_path: str | Path) -> list[_Section]:
    """DOCX를 (heading_path, 본문) 섹션 목록으로 — Heading 스타일이 경계.

    헤딩 텍스트는 본문에 넣지 않는다. heading_path가 들고 있고, 인덱스 입력엔
    rag/index_text가 앞에 붙이므로 본문에까지 넣으면 같은 문구가 두 번 들어간다.

    본문이 없는 헤딩(상위 목차처럼 바로 하위 헤딩이 오는 경우)은 섹션을 만들지 않는다 —
    '헤딩만 있고 내용 없는 청크'가 인덱스에 새는 것을 막는다(md 경로에서 실제로 관측된 노이즈).
    그 헤딩은 하위 섹션의 heading_path에 조상으로 남으므로 정보 손실은 없다.
    """
    sections: list[_Section] = []
    stack: list[str] = []      # 현재 헤딩 경로 — 인덱스가 곧 레벨-1
    lines: list[str] = []

    def flush() -> None:
        if lines:
            sections.append(_Section(list(stack), None, '\n'.join(lines)))   # page는 docx에 없음
            lines.clear()

    for is_heading, level, text in _docx_items(file_path):
        if is_heading:
            flush()                    # 이전 섹션 마감 후 경로 갱신
            del stack[level - 1:]      # 같은 레벨·하위 레벨 걷어내고
            stack.append(text)
        else:
            lines.append(text)
    flush()
    return sections


def _common_prefix(paths: list[list[str]]) -> list[str]:
    """heading_path들의 공통 조상 경로."""
    common = list(paths[0])
    for path in paths[1:]:
        while common != path[:len(common)]:
            common.pop()
    return common


def _pack_sections(
        sections: list[_Section],
        max_chars: int = _PACK_CHARS,
) -> list[_Section]:
    """연속 섹션을 예산까지 묶는다 — 섹션 하나가 청크 하나면 너무 잘아지므로.
    묶음의 page는 첫 섹션 것을 쓴다.

    묶음의 heading_path는 참여 섹션들의 **공통 조상**이고, 그보다 깊은 헤딩은 본문에
    줄로 남긴다. 정보 손실 없이 청크 크기만 옛 방식 수준으로 되돌린다.
      예: '2.1 신청 채널'과 '2.2 최소 후원 금액'을 묶으면
          heading_path=[문서명, '2. 정기후원 신청'], 본문 앞에 각 소제목이 한 줄씩.

    **묶기를 실제로 막는 건 예산이 아니라 경계다.** `joinable`이 공통 조상 깊이 2를
    요구하는데 heading_path[0]이 문서명이므로, 이는 "같은 대분류(H2) 아래 형제끼리만
    묶는다"는 뜻이다. 상위가 다르면 크기가 남아도 끊는다 — 무작정 묶으면 공통 조상이
    문서명 하나로 떨어져 heading_path가 무의미해지기 때문이다.
    그래서 실측 청크 중앙값은 max_chars(500)보다 훨씬 작은 200자대다 (2026-08-15).

    `size`는 body 길이만 센다 — flush가 덧붙이는 헤딩 줄은 빠지므로 max_chars는
    엄밀한 상한이 아니라 소프트 예산이다. 실제 상한은 뒤이어 SentenceSplitter가 건다.
    """
    packed: list[_Section] = []
    buf: list[_Section] = []
    size = 0

    def joinable(path: list[str]) -> bool:
        """buf에 이 섹션을 더해도 공통 조상이 문서명 아래로 유지되는가."""
        return len(_common_prefix([s.heading_path for s in buf] + [path])) >= 2

    def flush() -> None:
        nonlocal size
        if not buf:
            return
        common = _common_prefix([s.heading_path for s in buf])
        lines: list[str] = []
        for s in buf:
            lines.extend(s.heading_path[len(common):])   # 공통 조상보다 깊은 헤딩은 본문에 보존
            lines.append(s.body)
        packed.append(_Section(common, buf[0].page, '\n'.join(lines)))
        buf.clear()
        size = 0

    for section in sections:
        if buf and (size + len(section.body) > max_chars or not joinable(section.heading_path)):
            flush()
        buf.append(section)
        size += len(section.body)
    flush()
    return packed


def extract_text(file_path: str | Path) -> str:
    """파일 전체를 텍스트 하나로 추출한다 (청킹 없음). 채팅 첨부(컨텍스트 직접 주입)용.

    비ML 추출 — PDF 내 이미지·표·레이아웃 서식은 복원하지 않고 텍스트만 (채널톡 전략과 동일).
    """
    low = str(file_path).lower()
    if low.endswith('.pdf'):
        return '\n'.join(t for _, t in _pdf_pages(file_path)).strip()
    if low.endswith('.docx'):
        return _docx_text(file_path).strip()
    if low.endswith('.xlsx'):
        from rag.xlsx_chunking import chunk_xlsx     # 첨부 xlsx 지원 유지
        return '\n'.join(c.text for c in chunk_xlsx(str(file_path))).strip()
    return _read_text(file_path).strip()  # txt/md/기타


def _txt_sections(file_path: str | Path) -> list[_Section]:
    """평문 txt — 헤딩 개념이 없으므로 파일 전체가 섹션 하나다.

    빈 파일은 섹션 0개로 떨어뜨린다. split_text('')가 ['']를 반환해 빈 청크가
    새는 것을 막는다 (다른 파서의 `if buf:` 가드와 같은 목적).
    """
    text = _read_text(file_path)
    return [_Section([], None, text)] if text.strip() else []


def _md_sections_legacy(file_path: str | Path, splitter) -> list[ChunkData]:
    """md 기존 경로 — MarkdownNodeParser. 다음 커밋에서 _md_sections로 교체된다.

    이 함수만 `_Section` 파이프라인 밖에 있다. 교체가 동작 변경(청크 경계·본문이
    달라짐)이라 커밋을 분리했기 때문이고, 한 커밋 동안만 존재하는 과도기 형태다.
    """
    text = _read_text(file_path)
    nodes = splitter(MarkdownNodeParser()([LiDocument(text=text)]))
    return [
        ChunkData(
            text=node.get_content(),
            heading_path=_extract_heading_path(node.metadata, node.get_content()),
            page=None,
            chunk_index=i,
        )
        for i, node in enumerate(nodes)
    ]


# 형식 → 섹션 파서. 이 넷은 (heading_path, page, body)로 수렴하므로 뒤 단계를 공유한다.
# xlsx만 빠져 있다 — 시트=1청크·분할 없음·행상한 초과 시 거절이라 계약이 근본적으로 다르다.
_SECTION_PARSERS = {
    '.pdf': _pdf_sections,      # 글자 크기 휴리스틱
    '.docx': _docx_sections,    # Heading 스타일
    '.txt': _txt_sections,      # 헤딩 없음 — 통째로 섹션 1개
}


def chunk_file(file_path: str | Path, *, description: str = '') -> list[ChunkData]:
    """파일 한 개를 청크 리스트로. **형식 분기는 여기 한 곳뿐이다.**

    분기가 두 곳(여기 + 호출부)에 있던 게 CLI로 xlsx를 넣으면 ZIP 바이너리가 색인되던
    버그의 원인이었다(#42). 호출부는 이제 확장자를 보지 않는다.

    xlsx는 `chunk_xlsx`로 위임한다 — 시트=1청크, 분할 없음, 150행 초과 시 업로드 거절,
    `description` 병합. 나머지는 섹션 추출 → `_pack_sections` → 크기 분할을 공유한다.
    `description`은 xlsx 전용(다른 형식은 무시).
    """
    suffix = Path(file_path).suffix.lower()
    splitter = SentenceSplitter(chunk_size=512, chunk_overlap=50)

    if suffix == '.xlsx':
        from rag.xlsx_chunking import chunk_xlsx   # 지연 import (xlsx_chunking → chunking 순환 방지)
        return chunk_xlsx(str(file_path), description=description)

    if suffix == '.md':
        return _md_sections_legacy(file_path, splitter)

    parser = _SECTION_PARSERS.get(suffix)
    if parser is None:
        # 업로드는 SUPPORTED_SUFFIXES가 앞에서 거르므로 도달하지 않는다. 도달했다면
        # 지원 안 되는 파일이 파서를 잘못 타고 있다는 뜻이라 조용히 넘기지 않는다.
        raise ValueError(f'지원하지 않는 형식: {suffix or "확장자 없음"}')

    out: list[ChunkData] = []
    for section in _pack_sections(parser(file_path)):
        for chunk in splitter.split_text(section.body):
            out.append(ChunkData(text=chunk, heading_path=list(section.heading_path),
                                 page=section.page, chunk_index=len(out)))
    return out


def chunk_txt(file_path: str | Path) -> list[ChunkData]:
    """평문 txt 청킹 — `chunk_file`의 별칭 (기존 호출부·테스트 호환)."""
    return chunk_file(file_path)
