"""고난도 코퍼스 변환 스크립트 (#95) — _src md/json 원본 → 업로드 포맷 (재현 가능).

기존 corpus_v2는 md→PDF/DOCX 변환이 수동(비커밋)이라 재현이 안 됐다 — 이번 신규
문서(각 테넌트 11~13번)부터는 이 스크립트가 변환의 정본이다. 원본을 고치면 재실행으로
산출물이 따라온다.

파서 계약(rag/chunking.py·rag/xlsx_chunking.py)을 만족하도록 변환한다:
- PDF: 헤딩 판정이 **글자 크기**다(본문 최빈 크기보다 큰 줄) — # 레벨별로 본문(10pt)보다
  큰 16/13/11.5pt를 적용한다. 한글은 reportlab 내장 CID 폰트(HYGothic-Medium)로 —
  TTF 파일 의존이 없어 어느 환경에서든 같은 산출물이 나온다.
  md 표는 PDF에서 평문 줄로 렌더된다 — 의도된 것("표 모양은 유지 안 됨" 현실 재현).
- DOCX: 헤딩 판정이 **스타일명**(Heading N)이다 — add_heading(level=N)을 쓴다.
  md 표는 실제 워드 표로 변환한다(파서의 표 경로 시험).
- XLSX: 시트=청크 계약 — json 스펙({sheets: [{name, rows: [[...]]}]})에서 생성.
  첫 행 헤더·A1 시작·수치는 숫자 타입·시트당 150행 이하(스크립트가 검증).

사용: python sample_docs/corpus_v2/_src/build_hard_docs.py
  _src/<tenant>/<tenant>_1[123]_*.md|.xlsx.json 을 찾아 sample_docs/corpus_v2/<tenant>/ 에 산출.
  (기존 01~10 문서는 대상 아님 — 파일명 _1[123]_ 패턴만.)
"""
import json
import re
import sys
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

SRC = Path(__file__).resolve().parent
OUT_ROOT = SRC.parent
FONT = "HYGothic-Medium"
XLSX_MAX_ROWS = 150   # rag/xlsx_chunking.py와 같은 값 — 초과 산출물은 업로드가 거절된다

# 헤딩 크기 사다리 — 본문(10)보다 커야 _pdf_sections가 헤딩으로 인식한다.
PDF_SIZES = {0: 10, 1: 16, 2: 13, 3: 11.5}
_HEADING = re.compile(r"^(#{1,3})\s+(.*)")
_TABLE_ROW = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|\s*$")


def _md_blocks(text: str):
    """(level, content) 시퀀스 — level 0=본문 줄, 1~3=헤딩, 'table'=표 행 목록."""
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if m := _HEADING.match(line):
            yield len(m.group(1)), m.group(2).strip()
        elif _TABLE_ROW.match(line):
            rows = []
            while i < len(lines) and _TABLE_ROW.match(lines[i]):
                if not _TABLE_SEP.match(lines[i]):
                    rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            yield "table", rows
            continue
        elif line.strip():
            yield 0, line.strip()
        i += 1


def build_pdf(md_path: Path, out_path: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont(FONT))
    styles = {lv: ParagraphStyle(f"h{lv}", fontName=FONT, fontSize=sz, leading=sz * 1.5)
              for lv, sz in PDF_SIZES.items()}
    story = []
    for level, content in _md_blocks(md_path.read_text()):
        if level == "table":
            # 표를 평문 줄로 — PDF 표 모양이 안 살아남는 실서비스 조건의 정직한 재현
            for row in content:
                story.append(Paragraph(" | ".join(row), styles[0]))
        else:
            story.append(Paragraph(content, styles[level]))
            if level:
                story.append(Spacer(1, 6))
    SimpleDocTemplate(str(out_path), pagesize=A4).build(story)


def build_docx(md_path: Path, out_path: Path) -> None:
    doc = DocxDocument()
    for level, content in _md_blocks(md_path.read_text()):
        if level == "table":
            table = doc.add_table(rows=len(content), cols=len(content[0]))
            for r, row in enumerate(content):
                for c, cell in enumerate(row):
                    table.cell(r, c).text = cell
        elif level:
            doc.add_heading(content, level=level)   # 스타일명 'Heading N' — 파서 계약
        else:
            doc.add_paragraph(content)
    doc.save(str(out_path))


def build_xlsx(spec_path: Path, out_path: Path) -> None:
    spec = json.loads(spec_path.read_text())
    wb = Workbook()
    wb.remove(wb.active)
    for sheet in spec["sheets"]:
        rows = sheet["rows"]
        if len(rows) - 1 > XLSX_MAX_ROWS:
            raise SystemExit(f"{spec_path.name}/{sheet['name']}: 데이터 {len(rows)-1}행 > {XLSX_MAX_ROWS} — 업로드가 거절된다")
        ws = wb.create_sheet(sheet["name"])
        for row in rows:
            # 숫자 문자열은 숫자 타입으로 — xlsx_chunking·gold 오염검사가 기대하는 형태
            ws.append([float(c) if isinstance(c, str) and re.fullmatch(r"-?\d+(\.\d+)?", c)
                       else c for c in row])
    wb.save(str(out_path))


def main() -> None:
    built = []
    for tenant_dir in sorted(d for d in SRC.iterdir() if d.is_dir()):
        out_dir = OUT_ROOT / tenant_dir.name
        if not out_dir.is_dir():
            continue
        for src in sorted(tenant_dir.glob(f"{tenant_dir.name}_1[123]_*")):
            stem_target = src.name.removesuffix(".md").removesuffix(".xlsx.json")
            if src.name.endswith(".xlsx.json"):
                out = out_dir / f"{stem_target}.xlsx"
                build_xlsx(src, out)
            elif src.suffix == ".md":
                # 산출 형식은 파일명에 인코딩 — <이름>.pdf.md / <이름>.docx.md
                if src.name.endswith(".pdf.md"):
                    out = out_dir / src.name.removesuffix(".pdf.md").__add__(".pdf")
                    build_pdf(src, out)
                elif src.name.endswith(".docx.md"):
                    out = out_dir / src.name.removesuffix(".docx.md").__add__(".docx")
                    build_docx(src, out)
                else:
                    continue   # 확장 규약 밖 md(참고용)는 건드리지 않는다
            else:
                continue
            built.append(out)
            print(f"built {out.relative_to(OUT_ROOT)}")
    if not built:
        sys.exit("변환 대상 0건 — _src/<tenant>/<tenant>_1[123]_*.{pdf.md,docx.md,xlsx.json} 명명 확인")


if __name__ == "__main__":
    main()
